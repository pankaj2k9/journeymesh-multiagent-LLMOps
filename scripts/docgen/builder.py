"""Word-document building blocks for the JourneyMesh architecture guide.

Thin, opinionated helpers over python-docx so the content modules can read as
prose rather than as XML plumbing: headings, tables with captions, monospaced
code and diagram blocks, callouts, and the page furniture (cover, table of
contents field, running header, page numbers).
"""

from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from typing import Iterable, Optional, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# Palette - the JourneyMesh brand, muted for print
# ---------------------------------------------------------------------------
BRAND = RGBColor(0x17, 0x36, 0x5D)
ACCENT = RGBColor(0x1C, 0x63, 0xDC)
INK = RGBColor(0x12, 0x21, 0x2F)
MUTED = RGBColor(0x54, 0x64, 0x76)
CODE_BG = "F4F6F8"
DIAGRAM_BG = "F7F9FB"
CALLOUT_BG = {
    "note": "EEF6FF",
    "warning": "FFF8E6",
    "important": "FDEEF1",
    "tip": "ECFDF5",
}
CALLOUT_LABEL = {
    "note": "Note",
    "warning": "Caution",
    "important": "Important",
    "tip": "In practice",
}

MONO = "Consolas"
BODY = "Calibri"


def _shade(element, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    element.append(shading)


def _borders(element, colour: str = "D8DEE6", size: int = 6) -> None:
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        line = OxmlElement(f"w:{edge}")
        line.set(qn("w:val"), "single")
        line.set(qn("w:sz"), str(size))
        line.set(qn("w:space"), "6")
        line.set(qn("w:color"), colour)
        borders.append(line)
    element.append(borders)


def _field(paragraph, instruction: str, placeholder: str = "") -> None:
    """Insert a Word field (page numbers, TOC) that Word evaluates on open."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)

    if placeholder:
        text = OxmlElement("w:t")
        text.text = placeholder
        run._r.append(text)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


@dataclass
class DocumentMeta:
    title: str
    subtitle: str
    project: str
    tagline: str
    author: str
    email: str
    website: str
    version: str


class Guide:
    """The document under construction."""

    def __init__(self, meta: DocumentMeta) -> None:
        self.meta = meta
        self.doc = Document()
        self.table_count = 0
        self.figure_count = 0
        self._configure_styles()
        self._configure_page()

    # -- setup ------------------------------------------------------------
    def _configure_styles(self) -> None:
        styles = self.doc.styles

        normal = styles["Normal"]
        normal.font.name = BODY
        normal.font.size = Pt(10.5)
        normal.font.color.rgb = INK
        normal.paragraph_format.space_after = Pt(7)
        normal.paragraph_format.line_spacing = 1.15

        sizes = {"Heading 1": 20, "Heading 2": 15, "Heading 3": 12.5, "Heading 4": 11}
        for name, size in sizes.items():
            style = styles[name]
            style.font.name = BODY
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = BRAND if name == "Heading 1" else ACCENT
            style.paragraph_format.space_before = Pt(16 if name == "Heading 1" else 12)
            style.paragraph_format.space_after = Pt(6)
            style.paragraph_format.keep_with_next = True

        title = styles["Title"]
        title.font.name = BODY
        title.font.size = Pt(34)
        title.font.color.rgb = BRAND

    def _configure_page(self) -> None:
        section = self.doc.sections[0]
        section.page_width = Inches(8.27)   # A4
        section.page_height = Inches(11.69)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)

    def add_page_furniture(self) -> None:
        """Running header and a footer carrying the page number."""
        for section in self.doc.sections:
            section.different_first_page_header_footer = True

            header = section.header.paragraphs[0]
            header.text = f"{self.meta.project} - {self.meta.title.splitlines()[-1]}"
            header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in header.runs:
                run.font.size = Pt(8)
                run.font.color.rgb = MUTED

            footer = section.footer.paragraphs[0]
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = footer.add_run(f"{self.meta.author}  |  ")
            run.font.size = Pt(8)
            run.font.color.rgb = MUTED
            _field(footer, " PAGE ", "1")
            run = footer.add_run(" of ")
            run.font.size = Pt(8)
            run.font.color.rgb = MUTED
            _field(footer, " NUMPAGES ", "1")
            for run in footer.runs:
                run.font.size = Pt(8)
                run.font.color.rgb = MUTED

    # -- structure ---------------------------------------------------------
    def cover(self) -> None:
        meta = self.meta
        for _ in range(4):
            self.doc.add_paragraph()

        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(meta.project)
        run.font.size = Pt(44)
        run.font.bold = True
        run.font.color.rgb = BRAND

        sub = self.doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.add_run("Architecture Explanation Guide")
        run.font.size = Pt(24)
        run.font.color.rgb = ACCENT

        tag = self.doc.add_paragraph()
        tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = tag.add_run(meta.tagline)
        run.font.size = Pt(12)
        run.font.italic = True
        run.font.color.rgb = MUTED

        self.doc.add_paragraph()
        blurb = self.doc.add_paragraph()
        blurb.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = blurb.add_run(meta.subtitle)
        run.font.size = Pt(11.5)
        run.font.color.rgb = INK

        for _ in range(6):
            self.doc.add_paragraph()

        table = self.doc.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for label, value in (
            ("Author", meta.author),
            ("Email", meta.email),
            ("Website", meta.website),
            ("Version", meta.version),
            ("Generated", date.today().isoformat()),
        ):
            row = table.add_row().cells
            left = row[0].paragraphs[0]
            run = left.add_run(label)
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = MUTED
            left.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            right = row[1].paragraphs[0]
            run = right.add_run(value)
            run.font.size = Pt(10)
        for row in table.rows:
            row.cells[0].width = Inches(1.4)
            row.cells[1].width = Inches(3.2)

        self.page_break()

    def table_of_contents(self) -> None:
        self.h1("Table of Contents")
        note = self.doc.add_paragraph()
        run = note.add_run(
            "Word builds this from the document's headings. If it appears empty, "
            "right-click it and choose Update Field, or press F9."
        )
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = MUTED

        paragraph = self.doc.add_paragraph()
        _field(paragraph, r' TOC \o "1-3" \h \z \u ', "Right-click and choose Update Field.")
        self.page_break()

    def page_break(self) -> None:
        self.doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # -- headings ----------------------------------------------------------
    def h1(self, text: str, *, page_break: bool = False) -> None:
        if page_break:
            self.page_break()
        self.doc.add_heading(text, level=1)

    def h2(self, text: str) -> None:
        self.doc.add_heading(text, level=2)

    def h3(self, text: str) -> None:
        self.doc.add_heading(text, level=3)

    def h4(self, text: str) -> None:
        self.doc.add_heading(text, level=4)

    # -- text --------------------------------------------------------------
    def p(self, text: str, *, italic: bool = False, bold: bool = False) -> None:
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(text)
        run.font.italic = italic
        run.font.bold = bold

    def bullets(self, items: Iterable[str]) -> None:
        for item in items:
            self.doc.add_paragraph(item, style="List Bullet")

    def numbered(self, items: Iterable[str]) -> None:
        for item in items:
            self.doc.add_paragraph(item, style="List Number")

    def definition(self, term: str, technical: str, simple: str) -> None:
        """The two-level explanation used throughout the guide."""
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(f"{term} - technical. ")
        run.font.bold = True
        paragraph.add_run(technical)

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(f"{term} - in plain words. ")
        run.font.bold = True
        run.font.color.rgb = ACCENT
        paragraph.add_run(simple)

    def callout(self, kind: str, text: str) -> None:
        table = self.doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        _shade(cell._tc.get_or_add_tcPr(), CALLOUT_BG.get(kind, CALLOUT_BG["note"]))
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(f"{CALLOUT_LABEL.get(kind, 'Note')}. ")
        run.font.bold = True
        run.font.size = Pt(10)
        run = paragraph.add_run(text)
        run.font.size = Pt(10)
        self.doc.add_paragraph()

    # -- code and diagrams -------------------------------------------------
    def code(self, text: str, *, caption: Optional[str] = None, size: float = 8.5) -> None:
        table = self.doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        _shade(cell._tc.get_or_add_tcPr(), CODE_BG)
        cell.text = ""
        for index, line in enumerate(text.strip("\n").split("\n")):
            paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            run = paragraph.add_run(line if line else " ")
            run.font.name = MONO
            run.font.size = Pt(size)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), MONO)
        if caption:
            self.caption(caption, kind="listing")
        else:
            self.doc.add_paragraph()

    def diagram(self, text: str, caption: str, *, size: float = 8.0) -> None:
        table = self.doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        _shade(cell._tc.get_or_add_tcPr(), DIAGRAM_BG)
        cell.text = ""
        for index, line in enumerate(text.strip("\n").split("\n")):
            paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            run = paragraph.add_run(line if line else " ")
            run.font.name = MONO
            run.font.size = Pt(size)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), MONO)
        self.figure_count += 1
        self.caption(f"Figure {self.figure_count}. {caption}", kind="figure")

    def caption(self, text: str, kind: str = "figure") -> None:
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(text)
        run.font.size = Pt(8.5)
        run.font.italic = True
        run.font.color.rgb = MUTED

    # -- tables ------------------------------------------------------------
    def table(
        self,
        headers: Sequence[str],
        rows: Sequence[Sequence[str]],
        *,
        caption: Optional[str] = None,
        widths: Optional[Sequence[float]] = None,
        size: float = 9.0,
    ) -> None:
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        for index, header in enumerate(headers):
            cell = table.rows[0].cells[index]
            _shade(cell._tc.get_or_add_tcPr(), "E8EEF6")
            paragraph = cell.paragraphs[0]
            paragraph.text = ""
            run = paragraph.add_run(header)
            run.font.bold = True
            run.font.size = Pt(size)
            run.font.color.rgb = BRAND

        for row_values in rows:
            cells = table.add_row().cells
            for index, value in enumerate(row_values[: len(headers)]):
                paragraph = cells[index].paragraphs[0]
                paragraph.text = ""
                paragraph.paragraph_format.space_after = Pt(2)
                for line_index, line in enumerate(str(value).split("\n")):
                    target = paragraph if line_index == 0 else cells[index].add_paragraph()
                    run = target.add_run(line)
                    run.font.size = Pt(size)
                    if line.startswith("`") and line.endswith("`"):
                        run.text = line.strip("`")
                        run.font.name = MONO

        if widths:
            total = sum(widths)
            available = 6.4
            for row in table.rows:
                for index, weight in enumerate(widths[: len(headers)]):
                    row.cells[index].width = Inches(available * weight / total)

        self.table_count += 1
        self.caption(
            f"Table {self.table_count}. {caption}" if caption else f"Table {self.table_count}.",
            kind="table",
        )

    def understand(self, items: Sequence[str]) -> None:
        """The 'What You Should Understand' checkpoint after a major chapter."""
        self.h3("What you should understand")
        table = self.doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        _shade(cell._tc.get_or_add_tcPr(), "F1F5F9")
        cell.text = ""
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run("After this chapter you should be able to explain:")
        run.font.bold = True
        run.font.size = Pt(10)
        for index, item in enumerate(items, start=1):
            entry = cell.add_paragraph()
            entry.paragraph_format.space_after = Pt(2)
            run = entry.add_run(f"{index}. {item}")
            run.font.size = Pt(10)
        self.doc.add_paragraph()

    def qa(self, question: str, short: str, detail: str, follow_up: str) -> None:
        """One interview question in the house format."""
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(question)
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = BRAND

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run("Short answer. ")
        run.font.bold = True
        paragraph.add_run(short)

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run("Detail. ")
        run.font.bold = True
        paragraph.add_run(detail)

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run("Likely follow-up. ")
        run.font.bold = True
        run.font.italic = True
        run.font.color.rgb = MUTED
        run = paragraph.add_run(follow_up)
        run.font.italic = True
        run.font.color.rgb = MUTED

    def save(self, path: str) -> None:
        self.doc.save(path)
